"""
app.py — Legal Library Main App
⚖️ विधिक पुस्तकालय | UP Police Legal Reference System
"""

import os
import json
import streamlit as st
from pathlib import Path

from config import (
    APP_VERSION,
    BNS_JSON, CASE_LAWS_JSON, UPLOADS_DIR, AI_PROVIDERS
)

APP_TITLE    = "⚖️ विधिक पुस्तकालय"
APP_SUBTITLE = "UP Police Legal Reference System"
CCA_JSON     = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "cca_rules.json")
RTI_JSON     = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "rti_sections.json")

# ── Page config
st.set_page_config(
    page_title="विधिक पुस्तकालय",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Devanagari:wght@400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Noto Sans Devanagari', sans-serif; }
    .main-header { background: linear-gradient(135deg, #1a1d27, #21253a);
        padding: 1rem 1.5rem; border-radius: 10px; border-left: 4px solid #f4a623;
        margin-bottom: 1.5rem; }
    .main-header h1 { color: #f4a623; margin: 0; font-size: 1.6rem; }
    .main-header p  { color: #8890b5; margin: 0.2rem 0 0; font-size: 0.85rem; }
    .stat-card { background: #1a1d27; border: 1px solid #2e3250;
        border-radius: 8px; padding: 1rem; text-align: center; }
    .stat-card .num { font-size: 2rem; font-weight: 700; color: #4ecdc4; }
    .stat-card .lbl { font-size: 0.78rem; color: #8890b5; }
    .section-badge { background: #f4a623; color: #000; font-weight: 700;
        padding: 2px 8px; border-radius: 4px; font-size: 0.85rem; }
    .new-badge { background: #27ae60; color: #fff; font-size: 0.7rem;
        padding: 1px 5px; border-radius: 3px; margin-left: 4px; }
    .ai-provider { background: #21253a; border: 1px solid #2e3250;
        border-radius: 6px; padding: 0.4rem 0.8rem; font-size: 0.8rem;
        color: #4ecdc4; display: inline-block; margin-bottom: 0.5rem; }
    div[data-testid="stExpander"] { border: 1px solid #2e3250; border-radius: 8px; }
    .cca-card { background: #1a1d27; border: 1px solid #f4a623;
        border-radius: 8px; padding: 0.8rem 1rem; margin-bottom: 0.5rem; }
</style>
""", unsafe_allow_html=True)

# ── Helper functions
@st.cache_data
def load_bns_data():
    try:
        for p in [BNS_JSON, Path(__file__).parent/"data"/"bns_sections.json"]:
            if Path(p).exists():
                with open(p, encoding="utf-8") as f:
                    return json.load(f)
    except Exception:
        pass
    return []

@st.cache_data
def load_case_laws():
    try:
        for p in [CASE_LAWS_JSON, Path(__file__).parent/"data"/"case_laws.json"]:
            if Path(p).exists():
                with open(p, encoding="utf-8") as f:
                    return json.load(f)
    except Exception:
        pass
    return []

@st.cache_data
def load_rti_data():
    try:
        for p in [RTI_JSON, Path(__file__).parent/"data"/"rti_sections.json"]:
            if Path(p).exists():
                with open(p, encoding="utf-8") as f:
                    return json.load(f)
    except Exception:
        pass
    return []

@st.cache_data
def load_cca_data():
    try:
        for p in [CCA_JSON, Path(__file__).parent/"data"/"cca_rules.json"]:
            if Path(p).exists():
                with open(p, encoding="utf-8") as f:
                    return json.load(f)
    except Exception:
        pass
    return []

def save_case_laws(data: list):
    with open(CASE_LAWS_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    load_case_laws.clear()

def active_ai_providers():
    return [p for p in AI_PROVIDERS if p["key_env"] and not p["key_env"].startswith("your_")]

# ── Sidebar
with st.sidebar:
    st.markdown("## ⚖️ विधिक पुस्तकालय")
    st.markdown("*UP Police Legal Reference System*")
    st.divider()

    page = st.radio(
        "📌 Menu",
        options=[
            "🏠 Dashboard",
            "🔍 BNS धारा खोज",
            "📚 Case Law Library",
            "⚡ जमानत विरोध आख्या",
            "📝 RTI सहायक",
            "📋 विभागीय जाँच",
            "📥 Google Drive Sync",
            "⚙️ Settings & API Keys",
        ],
        label_visibility="collapsed"
    )

    st.divider()
    providers = active_ai_providers()
    if providers:
        st.markdown("**🤖 AI Providers Active:**")
        for p in providers:
            st.markdown(f"✅ {p['name']} — `{p['model']}`")
    else:
        st.warning("⚠️ कोई API key नहीं\nSettings में डालें")
    st.divider()
    st.caption(f"v{APP_VERSION} | UP Police Legal Ref.")

# ══════════════════════════════════════════════════════
# PAGE: Dashboard
# ══════════════════════════════════════════════════════
if page == "🏠 Dashboard":
    st.markdown("""
    <div class="main-header">
        <h1>⚖️ विधिक पुस्तकालय</h1>
        <p>UP Police Legal Reference System | BNS 2023 • RTI Act 2005 • CCA Rules 1991</p>
    </div>
    """, unsafe_allow_html=True)

    bns_data  = load_bns_data()
    case_laws = load_case_laws()
    rti_data  = load_rti_data()
    cca_data  = load_cca_data()

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""<div class="stat-card">
            <div class="num">{len(bns_data)}</div>
            <div class="lbl">BNS धाराएं</div></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="stat-card">
            <div class="num" style="color:#f4a623">{len(rti_data)}</div>
            <div class="lbl">RTI धाराएं</div></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="stat-card">
            <div class="num" style="color:#e05c5c">{len(cca_data)}</div>
            <div class="lbl">CCA Rules</div></div>""", unsafe_allow_html=True)
    with c4:
        ai_count = len(active_ai_providers())
        st.markdown(f"""<div class="stat-card">
            <div class="num" style="color:#27ae60">{ai_count}/4</div>
            <div class="lbl">AI Providers</div></div>""", unsafe_allow_html=True)

    st.divider()
    st.markdown("### ⚡ Quick Actions")
    qa1, qa2, qa3, qa4 = st.columns(4)
    with qa1:
        if st.button("🔍 BNS धारा खोजें", use_container_width=True):
            st.session_state["goto"] = "🔍 BNS धारा खोज"
            st.rerun()
    with qa2:
        if st.button("⚡ जमानत विरोध", use_container_width=True):
            st.session_state["goto"] = "⚡ जमानत विरोध आख्या"
            st.rerun()
    with qa3:
        if st.button("📝 RTI अपील", use_container_width=True):
            st.session_state["goto"] = "📝 RTI सहायक"
            st.rerun()
    with qa4:
        if st.button("📋 विभागीय जाँच", use_container_width=True):
            st.session_state["goto"] = "📋 विभागीय जाँच"
            st.rerun()

    st.divider()
    st.markdown("### 📋 Modules Overview")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
<div class="cca-card">
<b style="color:#f4a623">⚖️ BNS 2023</b><br>
<small style="color:#8890b5">भारतीय न्याय संहिता — 42 धाराएं, IPC mapping, जमानत विरोध आख्या</small>
</div>
<div class="cca-card">
<b style="color:#4ecdc4">📝 RTI Act 2005</b><br>
<small style="color:#8890b5">31 धाराएं, प्रथम अपील, द्वितीय अपील प्ररूप-14 (UP SIC)</small>
</div>
""", unsafe_allow_html=True)
    with col2:
        st.markdown("""
<div class="cca-card">
<b style="color:#e05c5c">📋 CCA Rules 1991</b><br>
<small style="color:#8890b5">विभागीय जाँच, आरोप पत्र जवाब, निलंबन, अपील, Revision</small>
</div>
<div class="cca-card">
<b style="color:#27ae60">📚 Case Law Library</b><br>
<small style="color:#8890b5">Manual + PDF से case laws जोड़ें, Google Drive sync</small>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# PAGE: BNS धारा खोज
# ══════════════════════════════════════════════════════
elif page == "🔍 BNS धारा खोज":
    st.header("🔍 BNS 2023 — धारा खोज")

    bns_data = load_bns_data()

    col1, col2 = st.columns([3, 2])
    with col1:
        query = st.text_input("🔍 BNS धारा / अपराध / हिंदी में खोजें",
                              placeholder="जैसे: हत्या, S.103, murder, 304...")
    with col2:
        ipc_query = st.text_input("📌 पुरानी IPC डालें → नई BNS देखें",
                                  placeholder="जैसे: 302, 376, 420...")

    f1, f2, f3 = st.columns(3)
    with f1:
        chunk_filter = st.multiselect("Chapter", options=[
            "जीवन अपराध", "चोट/हमला", "महिला/बच्चे",
            "राज्य के विरुद्ध", "संपत्ति अपराध", "जालसाजी/साक्ष्य", "लोक व्यवस्था"
        ])
    with f2:
        bail_filter = st.selectbox("जमानत", ["सभी", "जमानती", "गैर-जमानती"])
    with f3:
        only_new = st.checkbox("केवल नई धाराएं ★")

    if bns_data:
        results = bns_data
        q = query.strip().lower()
        iq = ipc_query.strip().lower().replace("ipc", "").strip()

        if q:
            results = [r for r in results if
                       q in r.get("bns","").lower() or
                       q in r.get("offence","").lower() or
                       q in r.get("offence_hi","").lower() or
                       q in r.get("punishment","").lower()]
        if iq:
            results = [r for r in results if iq in r.get("ipc","").lower().replace("ipc","").strip()]
        if chunk_filter:
            results = [r for r in results if r.get("chapter") in chunk_filter]
        if only_new:
            results = [r for r in results if r.get("is_new")]
        if bail_filter == "जमानती":
            results = [r for r in results if r.get("bailable")]
        elif bail_filter == "गैर-जमानती":
            results = [r for r in results if not r.get("bailable")]

        st.caption(f"🔎 {len(results)} धाराएं मिलीं")

        for row in results:
            with st.expander(
                f"{'🆕 ' if row.get('is_new') else ''}"
                f"**{row['bns']}** — {row.get('offence_hi', row.get('offence',''))}"
            ):
                c1, c2, c3 = st.columns(3)
                c1.metric("IPC (पुरानी)", row.get("ipc", "—"))
                c2.metric("सज़ा", row.get("punishment_short", "देखें"))
                c3.metric("जमानत", "✅ जमानती" if row.get("bailable") else "❌ गैर-जमानती")

                st.markdown(f"**सज़ा का विवरण:** {row.get('punishment','—')}")

                if row.get("case_laws"):
                    st.markdown("**📚 संबंधित Case Law:**")
                    for cl in row["case_laws"]:
                        st.markdown(f"- {cl}")

                btn_col1, btn_col2 = st.columns(2)
                with btn_col1:
                    if st.button(f"🤖 AI से विस्तार", key=f"exp_{row['bns']}"):
                        from modules.ai_engine import explain_section
                        with st.spinner("AI समझा रहा है..."):
                            res = explain_section(row["bns"], row.get("offence","") + " — " + row.get("punishment",""))
                        if res["success"]:
                            st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                            st.markdown(res["text"])
                        else:
                            st.error(f"AI Error: {res['error']}")
                with btn_col2:
                    if st.button(f"⚡ जमानत विरोध", key=f"bail_{row['bns']}"):
                        st.session_state["bail_section"] = row["bns"]
                        st.session_state["goto"] = "⚡ जमानत विरोध आख्या"
                        st.rerun()
    else:
        st.info("📂 BNS data load नहीं हुआ। Settings → Data Management से JSON load करें।")

# ══════════════════════════════════════════════════════
# PAGE: Case Law Library
# ══════════════════════════════════════════════════════
elif page == "📚 Case Law Library":
    st.header("📚 Case Law Library")

    case_laws = load_case_laws()

    tab1, tab2, tab3 = st.tabs(["📋 देखें", "➕ नया जोड़ें", "📄 PDF से Extract"])

    with tab1:
        search_cl = st.text_input("🔍 Case law खोजें", placeholder="धारा, case name, keyword...")
        results = case_laws
        if search_cl:
            q = search_cl.lower()
            results = [c for c in case_laws if
                       q in c.get("title","").lower() or
                       q in c.get("sections","").lower() or
                       q in c.get("ratio","").lower() or
                       q in c.get("citation","").lower()]

        st.caption(f"{len(results)} case laws")
        for cl in results:
            with st.expander(f"**{cl.get('title','—')}** | {cl.get('citation','—')}"):
                st.markdown(f"**धाराएं:** {cl.get('sections','—')}")
                st.markdown(f"**Ratio Decidendi:** {cl.get('ratio','—')}")
                if cl.get("prosecution_points"):
                    st.markdown("**अभियोजन/बचाव के लिए उपयोगी बिंदु:**")
                    for pt in cl["prosecution_points"]:
                        st.markdown(f"• {pt}")

    with tab2:
        st.markdown("### ➕ नया Case Law मैनुअल जोड़ें")
        with st.form("add_case_law"):
            t1, t2 = st.columns(2)
            title    = t1.text_input("Case Name *", placeholder="State vs. Ram Kumar 2019")
            citation = t2.text_input("Citation *", placeholder="2019 SCC 123 / 2019 (3) ADJ 45")
            sections = st.text_input("धाराएं *", placeholder="BNS S.103 / IPC 302 / CCA Rule 14 / RTI S.19")
            court    = st.text_input("Court", placeholder="Allahabad High Court / Supreme Court")
            category = st.selectbox("Category", ["BNS/IPC (आपराधिक)", "RTI", "विभागीय जाँच (CCA)", "अन्य"])
            ratio    = st.text_area("Ratio Decidendi (मुख्य निर्णय हिंदी में) *", height=100)
            pros_pts = st.text_area("उपयोगी बिंदु (एक लाइन = एक बिंदु)")
            submitted = st.form_submit_button("✅ Case Law जोड़ें", use_container_width=True)
            if submitted:
                if not (title and citation and sections and ratio):
                    st.error("⚠️ * वाले fields जरूरी हैं")
                else:
                    new_cl = {
                        "id": f"cl_{len(case_laws)+1:04d}",
                        "title": title, "citation": citation,
                        "sections": sections, "court": court,
                        "category": category, "ratio": ratio,
                        "prosecution_points": [p.strip() for p in pros_pts.split("\n") if p.strip()],
                    }
                    case_laws.append(new_cl)
                    save_case_laws(case_laws)
                    # Drive पर भी save करें
                    try:
                        from modules.drive_sync import DriveSync
                        ds = DriveSync()
                        if ds.is_connected():
                            ds.upload_json(CASE_LAWS_JSON, "case_laws.json")
                            st.success(f"✅ '{title}' जोड़ा गया और Drive पर save हुआ!")
                        else:
                            st.success(f"✅ '{title}' जोड़ा गया!")
                    except:
                        st.success(f"✅ '{title}' जोड़ा गया!")

    with tab3:
        st.markdown("### 📄 PDF/Text से Case Law Extract करें")
        uploaded = st.file_uploader("Judgment PDF या Text file upload करें", type=["pdf","txt"])
        if uploaded:
            raw_text = ""
            if uploaded.type == "application/pdf":
                try:
                    import pdfplumber
                    import io
                    with pdfplumber.open(io.BytesIO(uploaded.read())) as pdf:
                        raw_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
                except Exception as e:
                    st.error(f"PDF read error: {e}")
            else:
                raw_text = uploaded.read().decode("utf-8", errors="ignore")

            if raw_text:
                st.text_area("Extracted Text (preview)", raw_text[:1500], height=150)
                if st.button("🤖 AI से Case Law Extract करें"):
                    from modules.ai_engine import extract_case_law
                    with st.spinner("AI extract कर रहा है..."):
                        res = extract_case_law(raw_text)
                    if res["success"]:
                        st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                        st.code(res["text"], language="json")
                        st.info("ऊपर 'नया जोड़ें' tab में copy करके save करें")
                    else:
                        st.error(res["error"])

# ══════════════════════════════════════════════════════
# PAGE: जमानत विरोध आख्या
# ══════════════════════════════════════════════════════
elif page == "⚡ जमानत विरोध आख्या":
    st.header("⚡ जमानत विरोध आख्या")
    st.caption("जमानत प्रार्थनापत्र upload करें → AI बिंदुवार विरोध आख्या तैयार करेगा")

    if not active_ai_providers():
        st.error("⚠️ कोई AI API key नहीं है। Settings में डालें।")
        st.stop()

    col1, col2 = st.columns([1,1])

    with col1:
        st.markdown("### 📄 Step 1: जमानत प्रार्थनापत्र")
        bail_file = st.file_uploader("PDF/DOCX upload करें", type=["pdf","docx","txt"], key="bail_up")
        bail_text_manual = st.text_area("या सीधे text paste करें", height=200,
                                         placeholder="जमानत प्रार्थनापत्र का content यहाँ paste करें...")

    with col2:
        st.markdown("### 📝 Step 2: FIR व अभियोजन की जानकारी")
        fir_details = st.text_area("FIR नंबर, धारा, घटना का विवरण", height=120,
                                   placeholder="FIR No. 123/2024, थाना: कोतवाली, धारा: BNS S.103(1)...")
        extra_points = st.text_area("विरोध के अतिरिक्त बिंदु", height=120,
                                    placeholder="• आरोपी फरार रहा है\n• पीड़ित परिवार को धमकी\n• गवाह प्रभावित हो सकते हैं...")

    bail_text = bail_text_manual
    if bail_file and not bail_text_manual:
        try:
            if bail_file.type == "application/pdf":
                import pdfplumber, io
                with pdfplumber.open(io.BytesIO(bail_file.read())) as pdf:
                    bail_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            elif "word" in bail_file.type or bail_file.name.endswith(".docx"):
                from docx import Document
                import io
                doc = Document(io.BytesIO(bail_file.read()))
                bail_text = "\n".join(p.text for p in doc.paragraphs)
            else:
                bail_text = bail_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"File read error: {e}")

    st.divider()

    if "bail_section" in st.session_state:
        st.info(f"धारा {st.session_state['bail_section']} के लिए आख्या")

    generate_btn = st.button("⚡ जमानत विरोध आख्या तैयार करें",
                             use_container_width=True, type="primary")

    if generate_btn:
        if not bail_text and not fir_details:
            st.error("⚠️ जमानत प्रार्थनापत्र या FIR details जरूरी हैं")
        else:
            with st.spinner("🤖 AI आख्या तैयार कर रहा है..."):
                from modules.ai_engine import generate_bail_opposition
                res = generate_bail_opposition(
                    bail_text or "प्रार्थनापत्र उपलब्ध नहीं",
                    fir_details,
                    extra_points
                )

            if res["success"]:
                st.markdown(f'<div class="ai-provider">🤖 {res["provider"]} द्वारा तैयार</div>',
                            unsafe_allow_html=True)
                st.markdown("---")
                akhya_text = res["text"]
                st.markdown(akhya_text)
                st.markdown("---")
                try:
                    from modules.docx_export import create_akhya_docx
                    import io
                    docx_bytes = create_akhya_docx(akhya_text, fir_details)
                    st.download_button(
                        label="📥 DOCX Download करें",
                        data=docx_bytes,
                        file_name="jamant_virodh_akhya.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.warning(f"DOCX export error: {e}")
            else:
                st.error(f"❌ AI Error: {res['error']}")

# ══════════════════════════════════════════════════════
# PAGE: RTI सहायक
# ══════════════════════════════════════════════════════
elif page == "📝 RTI सहायक":
    st.header("📝 RTI सहायक — सूचना का अधिकार अधिनियम 2005")

    rti_data = load_rti_data()

    tab1, tab2, tab3 = st.tabs(["🔍 धारा खोजें", "✍️ RTI आवेदन बनाएं", "📤 अपील बनाएं"])

    with tab1:
        q = st.text_input("🔍 खोजें", placeholder="जैसे: अपील, दंड, 30 दिन, छूट...")
        only_imp = st.checkbox("केवल महत्वपूर्ण धाराएं")
        results = rti_data
        if q:
            results = [r for r in rti_data if
                q.lower() in r["title"].lower() or
                q.lower() in r["description"].lower() or
                any(q.lower() in s.lower() for s in r.get("subsections",[]))]
        if only_imp:
            results = [r for r in results if r.get("important")]
        st.caption(f"{len(results)} धाराएं मिलीं")
        for r in results:
            with st.expander(f"**धारा {r['section_no']}** — {r['title']} {'⭐' if r.get('important') else ''}"):
                st.markdown(f"*{r['chapter']}*")
                st.markdown(f"**सार:** {r['description']}")
                if r.get("clauses"):
                    for key, clause in r["clauses"].items():
                        if isinstance(clause, dict) and clause.get("details"):
                            title_key = clause.get("title", clause.get("term", key))
                            with st.expander(f"📌 {title_key}", expanded=False):
                                for d in clause["details"]:
                                    st.markdown(f"• {d}")
                else:
                    st.markdown("**विवरण:**")
                    for s in r.get("subsections",[]):
                        st.markdown(f"• {s}")
                if st.button(f"🤖 AI से विस्तार", key=f"rti_{r['section_no']}"):
                    from modules.ai_engine import ask_ai
                    with st.spinner("AI समझा रहा है..."):
                        res = ask_ai(
                            "आप RTI विशेषज्ञ हैं। धारा को UP Police के संदर्भ में सरल हिंदी में समझाएं।",
                            f"धारा {r['section_no']} — {r['title']}: {r['description']}"
                        )
                    if res["success"]:
                        st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                        st.markdown(res["text"])

    with tab2:
        st.markdown("### ✍️ RTI आवेदन तैयार करें")
        c1, c2 = st.columns(2)
        with c1:
            applicant_name = st.text_input("आवेदक का नाम")
            applicant_address = st.text_area("पता", height=80)
            department = st.text_input("विभाग/कार्यालय", placeholder="जैसे: पुलिस अधीक्षक कार्यालय, श्रावस्ती")
        with c2:
            pio_name = st.text_input("लोक सूचना अधिकारी", placeholder="जैसे: जनपद पुलिस अधीक्षक")
            info_needed = st.text_area("मांगी गई सूचना का विवरण *", height=120,
                placeholder="जैसे:\n1. दिनांक xx से xx तक की ACR प्रतियाँ\n2. विभागीय जाँच के आदेश की प्रति")
            purpose = st.text_input("उद्देश्य (optional)")

        if st.button("⚡ RTI आवेदन तैयार करें", type="primary", use_container_width=True):
            if not (applicant_name and department and info_needed):
                st.error("⚠️ नाम, विभाग और सूचना का विवरण जरूरी है")
            else:
                from modules.ai_engine import ask_ai
                with st.spinner("AI आवेदन तैयार कर रहा है..."):
                    res = ask_ai(
                        """आप RTI विशेषज्ञ हैं। UP Police के लिए औपचारिक RTI आवेदन हिंदी में तैयार करें।
Format: सेवा में..., विषय:..., महोदय, अनुच्छेद 1,2,3... में सूचना मांगें, धारा 6(1) RTI Act 2005 का हवाला दें, अंत में प्रार्थना।""",
                        f"""आवेदक: {applicant_name}\nपता: {applicant_address}\nविभाग: {department}\nPIO: {pio_name}\nमांगी गई सूचना: {info_needed}\nउद्देश्य: {purpose}"""
                    )
                if res["success"]:
                    st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                    st.markdown("---")
                    st.markdown(res["text"])
                    try:
                        from modules.docx_export import create_akhya_docx
                        docx_bytes = create_akhya_docx(res["text"], f"RTI आवेदन — {department}")
                        st.download_button("📥 DOCX Download", data=docx_bytes,
                            file_name="rti_avedan.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except Exception as e:
                        st.warning(f"DOCX: {e}")

    with tab3:
        st.markdown("### 📤 अपील तैयार करें")
        appeal_type = st.radio("अपील का प्रकार",
            ["प्रथम अपील — धारा 19(1)", "द्वितीय अपील — प्ररूप 14 (UP SIC)"],
            horizontal=True)
        st.divider()

        if "प्रथम" in appeal_type:
            st.markdown("#### प्रथम अपील — धारा 19(1) RTI Act 2005")
            c1, c2 = st.columns(2)
            with c1:
                fa_name = st.text_input("आवेदक का नाम *", key="fa_name")
                fa_address = st.text_area("आवेदक का पता *", height=70, key="fa_addr")
                fa_dept = st.text_input("लोक प्राधिकरण/विभाग *", key="fa_dept")
                fa_pio = st.text_input("लोक सूचना अधिकारी", key="fa_pio")
            with c2:
                fa_rti_date = st.text_input("RTI आवेदन तारीख *", key="fa_date", placeholder="01/01/2025")
                fa_rti_no = st.text_input("RTI आवेदन संख्या", key="fa_no")
                fa_info = st.text_area("मांगी गई सूचना *", height=70, key="fa_info")
                fa_status = st.selectbox("क्या हुआ?", [
                    "30 दिन में कोई जवाब नहीं मिला",
                    "अधूरी सूचना दी गई",
                    "सूचना देने से मना किया गया",
                    "गलत/भ्रामक सूचना दी गई"
                ], key="fa_status")

            if st.button("⚡ प्रथम अपील तैयार करें", type="primary", use_container_width=True, key="fa_btn"):
                if not (fa_name and fa_dept and fa_info and fa_rti_date):
                    st.error("⚠️ * वाले fields जरूरी हैं")
                else:
                    from modules.ai_engine import ask_ai
                    with st.spinner("AI प्रथम अपील तैयार कर रहा है..."):
                        res = ask_ai(
                            "आप RTI Act 2005 के विशेषज्ञ हैं। धारा 19(1) के अंतर्गत प्रथम अपील औपचारिक हिंदी में तैयार करें। Format: सेवा में, विषय, महोदय, पृष्ठभूमि, अपील के आधार बिंदुवार (धारा 7(1)/7(2) का हवाला), प्रार्थना।",
                            f"आवेदक: {fa_name}\nपता: {fa_address}\nविभाग: {fa_dept}\nPIO: {fa_pio}\nRTI तारीख: {fa_rti_date}\nRTI संख्या: {fa_rti_no}\nसूचना: {fa_info}\nस्थिति: {fa_status}"
                        )
                    if res["success"]:
                        st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                        st.markdown("---")
                        st.markdown(res["text"])
                        try:
                            from modules.docx_export import create_akhya_docx
                            docx_bytes = create_akhya_docx(res["text"], f"प्रथम अपील — {fa_dept}")
                            st.download_button("📥 प्रथम अपील DOCX", data=docx_bytes,
                                file_name="rti_pratham_appeal.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True)
                        except Exception as e:
                            st.warning(f"DOCX: {e}")
                    else:
                        st.error(f"❌ {res['error']}")
        else:
            st.markdown("#### द्वितीय अपील — प्ररूप 14 | उत्तर प्रदेश राज्य सूचना आयोग")
            st.info("📋 UP SIC में द्वितीय अपील प्ररूप-14 में दाखिल होती है — धारा 19(3) RTI Act 2005")

            st.markdown("##### 🔵 आवेदक की जानकारी")
            c1, c2 = st.columns(2)
            with c1:
                sa_name = st.text_input("अपीलार्थी का नाम *", key="sa_name")
                sa_father = st.text_input("पिता/पति का नाम", key="sa_father")
                sa_address = st.text_area("पूरा पता *", height=70, key="sa_addr")
                sa_mobile = st.text_input("मोबाइल नंबर", key="sa_mobile")
            with c2:
                sa_dept = st.text_input("लोक प्राधिकरण *", key="sa_dept")
                sa_pio = st.text_input("लोक सूचना अधिकारी", key="sa_pio")
                sa_fao = st.text_input("प्रथम अपीलीय अधिकारी", key="sa_fao")
                sa_district = st.text_input("जनपद", key="sa_district")

            st.markdown("##### 🔵 RTI आवेदन का विवरण")
            c1, c2, c3 = st.columns(3)
            with c1:
                sa_rti_date = st.text_input("RTI आवेदन तारीख *", key="sa_rti_date", placeholder="दि. 01/01/2025")
                sa_rti_no = st.text_input("RTI क्रमांक", key="sa_rti_no")
            with c2:
                sa_fa_date = st.text_input("प्रथम अपील तारीख", key="sa_fa_date")
                sa_fa_no = st.text_input("प्रथम अपील क्रमांक", key="sa_fa_no")
            with c3:
                sa_today = st.text_input("आज की तारीख *", key="sa_today", placeholder="दि. 01/03/2025")
                sa_fee = st.text_input("शुल्क", key="sa_fee", placeholder="₹10 IPO/DD")

            sa_info = st.text_area("RTI में मांगी गई सूचना *", height=100, key="sa_info")

            st.markdown("##### 🔵 अपील के आधार")
            sg1 = st.checkbox("✅ 30 दिन में सूचना नहीं मिली — धारा 7(1)", key="sg1", value=True)
            sg2 = st.checkbox("✅ प्रथम अपील का निर्णय नहीं मिला — धारा 19(1)", key="sg2")
            sg3 = st.checkbox("✅ अधूरी/भ्रामक/गलत सूचना दी गई — धारा 18(1)(e)", key="sg3")
            sg4 = st.checkbox("✅ सूचना देने से अनुचित रूप से मना — धारा 7(8)", key="sg4")
            sa_extra = st.text_area("अन्य आधार/तथ्य", height=60, key="sa_extra")

            st.markdown("##### 🔵 प्रार्थना")
            sp1 = st.checkbox("✅ मांगी गई सूचना देने का निर्देश दिया जाए", key="sp1", value=True)
            sp2 = st.checkbox("✅ ₹250/दिन दंड अधिरोपित किया जाए — धारा 20(1)", key="sp2")
            sp3 = st.checkbox("✅ विभागीय कार्रवाई की संस्तुति — धारा 20(2)", key="sp3")
            sp4 = st.checkbox("✅ हानि का प्रतिकर दिलाया जाए — धारा 19(8)(b)", key="sp4")
            sp_extra = st.text_input("अन्य प्रार्थना", key="sp_extra")

            st.divider()
            if st.button("⚡ प्ररूप-14 द्वितीय अपील तैयार करें", type="primary", use_container_width=True, key="sa_btn"):
                if not (sa_name and sa_dept and sa_info and sa_rti_date and sa_today):
                    st.error("⚠️ * वाले fields जरूरी हैं")
                else:
                    grounds = []
                    if sg1: grounds.append("30 दिन की समय सीमा में सूचना नहीं मिली [धारा 7(1)]")
                    if sg2: grounds.append("प्रथम अपील का निर्णय नहीं मिला / असंतोषजनक [धारा 19(1)]")
                    if sg3: grounds.append("अधूरी/भ्रामक/गलत सूचना दी गई [धारा 18(1)(e)]")
                    if sg4: grounds.append("सूचना देने से अनुचित रूप से मना किया गया [धारा 7(8)]")
                    if sa_extra: grounds.append(sa_extra)
                    prayers = []
                    if sp1: prayers.append("मांगी गई सूचना देने का निर्देश दिया जाए")
                    if sp2: prayers.append("लोक सूचना अधिकारी पर ₹250 प्रतिदिन दंड अधिरोपित किया जाए [धारा 20(1)]")
                    if sp3: prayers.append("विभागीय कार्रवाई की संस्तुति की जाए [धारा 20(2)]")
                    if sp4: prayers.append("अपीलार्थी को हुई हानि का प्रतिकर दिलाया जाए [धारा 19(8)(b)]")
                    if sp_extra: prayers.append(sp_extra)

                    from modules.ai_engine import ask_ai
                    with st.spinner("AI प्ररूप-14 तैयार कर रहा है..."):
                        res = ask_ai(
                            """आप UP State Information Commission (UPSIC) के RTI विशेषज्ञ हैं। प्ररूप-14 के अनुसार द्वितीय अपील औपचारिक हिंदी में तैयार करें।
Format:
उत्तर प्रदेश राज्य सूचना आयोग के समक्ष
प्ररूप — 14
[सूचना का अधिकार अधिनियम, 2005 की धारा 19(3)]
द्वितीय अपील
अपीलार्थी: [नाम, पता]
विपक्षी: [विभाग, PIO]
1. RTI आवेदन का विवरण
2. प्रथम अपील का विवरण
3. अपील के आधार (बिंदुवार, धारा सहित)
4. प्रार्थना (बिंदुवार)
5. घोषणा
दिनांक/स्थान/हस्ताक्षर""",
                            f"""अपीलार्थी: {sa_name}, पिता/पति: {sa_father}
पता: {sa_address}, मोबाइल: {sa_mobile}, जनपद: {sa_district}
विभाग: {sa_dept}, PIO: {sa_pio}, FAO: {sa_fao}
RTI तारीख: {sa_rti_date}, RTI क्रमांक: {sa_rti_no}
प्रथम अपील तारीख: {sa_fa_date}, क्रमांक: {sa_fa_no}
शुल्क: {sa_fee}, आज की तारीख: {sa_today}
मांगी गई सूचना: {sa_info}
अपील के आधार: {chr(10).join(f'{i+1}. {g}' for i,g in enumerate(grounds))}
प्रार्थना: {chr(10).join(f'{i+1}. {p}' for i,p in enumerate(prayers))}"""
                        )
                    if res["success"]:
                        st.markdown(f'<div class="ai-provider">🤖 {res["provider"]} — प्ररूप 14</div>', unsafe_allow_html=True)
                        st.markdown("---")
                        appeal_text = res["text"]
                        st.markdown(appeal_text)
                        st.divider()
                        try:
                            from modules.docx_export import create_akhya_docx
                            docx_bytes = create_akhya_docx(appeal_text, f"द्वितीय अपील प्ररूप-14 | {sa_dept}")
                            st.download_button("📥 प्ररूप-14 DOCX Download", data=docx_bytes,
                                file_name=f"prarup14_{sa_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True)
                        except Exception as e:
                            st.warning(f"DOCX: {e}")
                    else:
                        st.error(f"❌ {res['error']}")

# ══════════════════════════════════════════════════════
# PAGE: विभागीय जाँच
# ══════════════════════════════════════════════════════
elif page == "📋 विभागीय जाँच":
    st.header("📋 विभागीय जाँच सहायक — CCA Rules 1991")
    st.caption("UP Police CCA Rules 1991 | आरोप पत्र का जवाब | अपील | Revision")

    cca_data = load_cca_data()

    tab1, tab2, tab3, tab4 = st.tabs([
        "📖 CCA Rules Reference",
        "📄 आरोप पत्र → जवाब",
        "⚖️ निलंबन/वेतन/बर्खास्तगी",
        "📤 अपील/Revision"
    ])

    # ── Tab 1: CCA Rules Reference
    with tab1:
        st.markdown("### 📖 UP Police CCA Rules 1991 — Reference")
        q_cca = st.text_input("🔍 Rule खोजें", placeholder="जैसे: निलंबन, बर्खास्तगी, जाँच, अपील...")
        only_imp_cca = st.checkbox("केवल महत्वपूर्ण Rules", value=True)

        results = cca_data
        if q_cca:
            results = [r for r in cca_data if
                q_cca.lower() in str(r.get("rule_no","")).lower() or
                q_cca.lower() in r.get("title","").lower() or
                q_cca.lower() in r.get("description","").lower() or
                any(q_cca.lower() in s.lower() for s in r.get("subsections",[]))]
        if only_imp_cca:
            results = [r for r in results if r.get("important")]

        st.caption(f"{len(results)} Rules मिले")
        for r in results:
            icon = "⚖️" if r.get("penalty_type") == "major" else "📋"
            with st.expander(f"{icon} **Rule {r['rule_no']}** — {r['title']}"):
                st.markdown(f"*{r['chapter']}*")
                st.markdown(f"**सार:** {r['description']}")
                if r.get("inquiry_required") is not None:
                    req = "✅ अनिवार्य" if r["inquiry_required"] else "❌ अनिवार्य नहीं"
                    st.markdown(f"**विस्तृत जाँच:** {req}")
                st.markdown("**विवरण:**")
                for s in r.get("subsections", []):
                    st.markdown(f"• {s}")
                if st.button(f"🤖 AI से विस्तार", key=f"cca_{r['rule_no']}"):
                    from modules.ai_engine import ask_ai
                    with st.spinner("AI समझा रहा है..."):
                        res = ask_ai(
                            "आप UP Police CCA Rules 1991 के विशेषज्ञ हैं। Rule को UP Police कर्मचारियों के लिए सरल हिंदी में समझाएं। Case laws का भी संदर्भ दें।",
                            f"Rule {r['rule_no']} — {r['title']}: {r['description']}\nविवरण: {'; '.join(r.get('subsections',[]))}"
                        )
                    if res["success"]:
                        st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                        st.markdown(res["text"])

    # ── Tab 2: आरोप पत्र → जवाब
    with tab2:
        st.markdown("### 📄 आरोप पत्र का जवाब (Rule 14)")
        st.info("आरोप पत्र upload करें या paste करें → AI बिंदुवार जवाब तैयार करेगा")

        c1, c2 = st.columns([1,1])
        with c1:
            st.markdown("**📄 आरोप पत्र**")
            charge_file = st.file_uploader("PDF/DOCX upload करें", type=["pdf","docx","txt"], key="charge_up")
            charge_text = st.text_area("या आरोपों का विवरण paste करें", height=200,
                placeholder="आरोप 1: दिनांक xx को अनुपस्थित रहे\nआरोप 2: उच्च अधिकारी के आदेश का पालन नहीं किया...")

        with c2:
            st.markdown("**📝 कर्मचारी की जानकारी**")
            emp_name = st.text_input("कर्मचारी का नाम", key="emp_name")
            emp_post = st.text_input("पद/बल संख्या", key="emp_post",
                placeholder="जैसे: आरक्षी, बल सं. 1234")
            emp_service = st.text_input("सेवा अवधि", key="emp_service",
                placeholder="जैसे: 15 वर्ष")
            emp_record = st.selectbox("सेवा रिकॉर्ड", [
                "स्वच्छ (कोई पूर्व दंड नहीं)",
                "सामान्य",
                "एक पूर्व लघु दंड"
            ], key="emp_record")
            emp_family = st.text_input("आश्रित परिवार", key="emp_family",
                placeholder="जैसे: पत्नी + 3 बच्चे, बीमार माता")
            emp_defense = st.text_area("बचाव के मुख्य बिंदु", height=100, key="emp_defense",
                placeholder="जैसे:\n• उस दिन अस्वस्थ था, चिकित्सक से परामर्श किया\n• आदेश की जानकारी नहीं थी\n• जाँच में प्रक्रिया का उल्लंघन हुआ है")

        # File extract
        if charge_file and not charge_text:
            try:
                if charge_file.type == "application/pdf":
                    import pdfplumber, io
                    with pdfplumber.open(io.BytesIO(charge_file.read())) as pdf:
                        charge_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
                elif charge_file.name.endswith(".docx"):
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(charge_file.read()))
                    charge_text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    charge_text = charge_file.read().decode("utf-8", errors="ignore")
            except Exception as e:
                st.error(f"File read error: {e}")

        st.divider()
        if st.button("⚡ आरोप पत्र का जवाब तैयार करें", type="primary", use_container_width=True, key="charge_btn"):
            if not charge_text:
                st.error("⚠️ आरोप पत्र का विवरण जरूरी है")
            else:
                from modules.ai_engine import ask_ai
                with st.spinner("AI जवाब तैयार कर रहा है..."):
                    res = ask_ai(
                        """आप UP Police CCA Rules 1991 के विशेषज्ञ हैं और कर्मचारी के बचाव पक्ष में हैं।
आरोप पत्र पढ़कर Rule 14 के अंतर्गत बिंदुवार जवाब हिंदी में तैयार करें।

Format:
सेवा में,
जाँच अधिकारी महोदय
[विभाग]

विषय: आरोप पत्र के संदर्भ में लिखित जवाब

महोदय,

प्रस्तर 1: परिचय (नाम, पद, सेवा अवधि)
प्रस्तर 2: प्रत्येक आरोप का बिंदुवार खंडन (आरोप 1 — खंडन, आरोप 2 — खंडन...)
प्रस्तर 3: प्राकृतिक न्याय के सिद्धांत (यदि प्रक्रिया में कोई त्रुटि हो)
प्रस्तर 4: सेवा रिकॉर्ड और परिस्थितियाँ
प्रस्तर 5: प्रार्थना — आरोपों से दोषमुक्त किया जाए

अपराध सिद्ध न होने पर दंड नहीं, प्रक्रिया का उल्लंघन जाँच को दोषपूर्ण बनाता है — Supreme Court के निर्णयों का हवाला दें।""",
                        f"""कर्मचारी: {emp_name}, पद: {emp_post}
सेवा अवधि: {emp_service}, रिकॉर्ड: {emp_record}
आश्रित परिवार: {emp_family}
बचाव के बिंदु: {emp_defense}

आरोप पत्र का विषय:
{charge_text[:4000]}"""
                    )
                if res["success"]:
                    st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                    st.markdown("---")
                    jawab_text = res["text"]
                    st.markdown(jawab_text)
                    st.divider()
                    try:
                        from modules.docx_export import create_akhya_docx
                        docx_bytes = create_akhya_docx(jawab_text, f"आरोप पत्र जवाब — {emp_name}")
                        st.download_button("📥 जवाब DOCX Download", data=docx_bytes,
                            file_name=f"aarop_jawab_{emp_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except Exception as e:
                        st.warning(f"DOCX: {e}")
                else:
                    st.error(f"❌ {res['error']}")

    # ── Tab 3: निलंबन/वेतन/बर्खास्तगी
    with tab3:
        st.markdown("### ⚖️ विशेष मामलों का जवाब/आवेदन")
        case_type = st.selectbox("मामले का प्रकार", [
            "निलंबन आदेश का विरोध (Rule 9)",
            "निर्वाह भत्ता बढ़ाने का आवेदन (Rule 9)",
            "वेतन वृद्धि रोकने का जवाब (Rule 23)",
            "अनिवार्य सेवानिवृत्ति का विरोध (Rule 8)",
            "बर्खास्तगी आदेश का विरोध (Rule 8)",
            "दंड कम करने का आवेदन"
        ], key="case_type")

        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            ct_name = st.text_input("कर्मचारी का नाम *", key="ct_name")
            ct_post = st.text_input("पद/बल संख्या", key="ct_post")
            ct_dept = st.text_input("विभाग/थाना", key="ct_dept")
            ct_service = st.text_input("सेवा अवधि", key="ct_service")
        with c2:
            ct_order_date = st.text_input("आदेश की तारीख", key="ct_order_date")
            ct_order_no = st.text_input("आदेश संख्या", key="ct_order_no")
            ct_authority = st.text_input("आदेश देने वाला प्राधिकारी", key="ct_authority",
                placeholder="जैसे: पुलिस अधीक्षक, श्रावस्ती")
            ct_family = st.text_input("आश्रित परिवार", key="ct_family")

        ct_facts = st.text_area("मामले के मुख्य तथ्य *", height=120, key="ct_facts",
            placeholder="जैसे:\n• किस कारण से निलंबित/दंडित किया गया\n• आरोप क्या था\n• कर्मचारी का पक्ष क्या है")

        if st.button("⚡ आवेदन/जवाब तैयार करें", type="primary", use_container_width=True, key="ct_btn"):
            if not (ct_name and ct_facts):
                st.error("⚠️ नाम और तथ्य जरूरी हैं")
            else:
                from modules.ai_engine import ask_ai
                with st.spinner("AI तैयार कर रहा है..."):
                    res = ask_ai(
                        f"""आप UP Police CCA Rules 1991 के विशेषज्ञ हैं और कर्मचारी के बचाव पक्ष में हैं।
{case_type} के लिए औपचारिक आवेदन/जवाब हिंदी में तैयार करें।

Format: सेवा में, विषय, महोदय, प्रस्तर 1 से शुरू करके बिंदुवार, CCA Rules और Supreme Court के निर्णयों का हवाला, प्रार्थना।

विशेष ध्यान:
- निलंबन: Ajay Kumar Choudhary 2015 SC का हवाला दें
- वेतन वृद्धि: Rule 7 में जाँच अनिवार्य नहीं पर कारण बताओ नोटिस आवश्यक है
- बर्खास्तगी: B.C. Chaturvedi 1995 SC — दंड अनुपात में होना चाहिए
- सेवा अवधि और स्वच्छ रिकॉर्ड को अवश्य उठाएं""",
                        f"""मामले का प्रकार: {case_type}
कर्मचारी: {ct_name}, पद: {ct_post}
विभाग: {ct_dept}, सेवा: {ct_service}
आदेश तारीख: {ct_order_date}, आदेश संख्या: {ct_order_no}
प्राधिकारी: {ct_authority}
आश्रित परिवार: {ct_family}
मुख्य तथ्य: {ct_facts}"""
                    )
                if res["success"]:
                    st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                    st.markdown("---")
                    ct_text = res["text"]
                    st.markdown(ct_text)
                    try:
                        from modules.docx_export import create_akhya_docx
                        docx_bytes = create_akhya_docx(ct_text, f"{case_type} — {ct_name}")
                        st.download_button("📥 DOCX Download", data=docx_bytes,
                            file_name=f"vibhagiya_{ct_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except Exception as e:
                        st.warning(f"DOCX: {e}")
                else:
                    st.error(f"❌ {res['error']}")

    # ── Tab 4: अपील/Revision
    with tab4:
        st.markdown("### 📤 विभागीय अपील / Revision")
        appeal_to = st.radio("अपील किसके समक्ष?", [
            "प्रथम अपील — DIG के समक्ष (Rule 16)",
            "द्वितीय अपील — IGP/ADG के समक्ष (Rule 16)",
            "Revision — DGP/शासन के समक्ष (Rule 17)"
        ], key="vib_appeal_to")

        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            va_name = st.text_input("कर्मचारी का नाम *", key="va_name")
            va_post = st.text_input("पद/बल संख्या", key="va_post")
            va_dept = st.text_input("वर्तमान तैनाती", key="va_dept")
            va_penalty = st.text_input("दिया गया दंड *", key="va_penalty",
                placeholder="जैसे: वेतन वृद्धि रोकी गई / निलंबित / बर्खास्त")
        with c2:
            va_order_date = st.text_input("दंड आदेश तारीख *", key="va_order_date")
            va_order_no = st.text_input("आदेश संख्या", key="va_order_no")
            va_service = st.text_input("सेवा अवधि", key="va_service")
            va_record = st.selectbox("सेवा रिकॉर्ड", [
                "स्वच्छ (कोई पूर्व दंड नहीं)",
                "सामान्य",
                "एक पूर्व लघु दंड"
            ], key="va_record")

        va_grounds = st.text_area("अपील के आधार *", height=120, key="va_grounds",
            placeholder="जैसे:\n• जाँच में प्राकृतिक न्याय का उल्लंघन हुआ\n• आरोप सिद्ध नहीं हुए\n• दंड अपराध के अनुपात में नहीं है\n• जाँच रिपोर्ट की प्रति नहीं दी गई")
        va_family = st.text_input("आश्रित परिवार", key="va_family")

        st.divider()
        if st.button("⚡ अपील तैयार करें", type="primary", use_container_width=True, key="va_btn"):
            if not (va_name and va_penalty and va_order_date and va_grounds):
                st.error("⚠️ * वाले fields जरूरी हैं")
            else:
                from modules.ai_engine import ask_ai
                with st.spinner("AI अपील तैयार कर रहा है..."):
                    res = ask_ai(
                        f"""आप UP Police CCA Rules 1991 के विशेषज्ञ हैं।
{appeal_to} के लिए औपचारिक अपील हिंदी में तैयार करें।

Format:
सेवा में, [अपीलीय प्राधिकारी]
विषय: CCA Rules 1991 के Rule 16/17 के अंतर्गत अपील
महोदय,
प्रस्तर 1: परिचय और पृष्ठभूमि
प्रस्तर 2: दंड आदेश का विवरण
प्रस्तर 3: अपील के आधार बिंदुवार (प्राकृतिक न्याय, Case Laws सहित)
प्रस्तर 4: सेवा रिकॉर्ड और व्यक्तिगत परिस्थितियाँ
प्रस्तर 5: प्रार्थना — दंड आदेश निरस्त/कम किया जाए

महत्वपूर्ण केस Laws:
- Ranjit Thakur vs Union of India 1987 SC — असमानुपातिक दंड
- B.C. Chaturvedi 1995 SC — दंड अनुपात में होना चाहिए
- Union of India vs Mohd. Ramzan Khan 1991 SC — जाँच रिपोर्ट प्रति अधिकार""",
                        f"""अपील: {appeal_to}
कर्मचारी: {va_name}, पद: {va_post}
तैनाती: {va_dept}, सेवा: {va_service}, रिकॉर्ड: {va_record}
दिया गया दंड: {va_penalty}
आदेश तारीख: {va_order_date}, आदेश संख्या: {va_order_no}
आश्रित परिवार: {va_family}
अपील के आधार: {va_grounds}"""
                    )
                if res["success"]:
                    st.markdown(f'<div class="ai-provider">🤖 {res["provider"]}</div>', unsafe_allow_html=True)
                    st.markdown("---")
                    appeal_text = res["text"]
                    st.markdown(appeal_text)
                    try:
                        from modules.docx_export import create_akhya_docx
                        docx_bytes = create_akhya_docx(appeal_text, f"विभागीय अपील — {va_name}")
                        st.download_button("📥 अपील DOCX Download", data=docx_bytes,
                            file_name=f"vibhagiya_appeal_{va_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except Exception as e:
                        st.warning(f"DOCX: {e}")
                else:
                    st.error(f"❌ {res['error']}")

# ══════════════════════════════════════════════════════
# PAGE: Google Drive Sync
# ══════════════════════════════════════════════════════
elif page == "📥 Google Drive Sync":
    st.header("📥 Google Drive Sync")

    try:
        from modules.drive_sync import DriveSync
        ds = DriveSync()
        if ds.is_connected():
            st.success("✅ Google Drive Connected")
            st.divider()

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("### 📤 Drive पर Upload")
                if st.button("📤 Case Laws Upload", use_container_width=True):
                    with st.spinner("Upload हो रहा है..."):
                        ok = ds.upload_json(CASE_LAWS_JSON, "case_laws.json")
                    st.success("✅ Upload complete!") if ok else st.error("❌ Upload failed")
                if st.button("📤 BNS Sections Upload", use_container_width=True):
                    with st.spinner("Upload हो रहा है..."):
                        ok = ds.upload_json(BNS_JSON, "bns_sections.json")
                    st.success("✅ Upload complete!") if ok else st.error("❌ Upload failed")
                if st.button("📤 RTI Sections Upload", use_container_width=True):
                    with st.spinner("Upload हो रहा है..."):
                        ok = ds.upload_json(RTI_JSON, "rti_sections.json")
                    st.success("✅ Upload complete!") if ok else st.error("❌ Upload failed")
                if st.button("📤 CCA Rules Upload", use_container_width=True):
                    with st.spinner("Upload हो रहा है..."):
                        ok = ds.upload_json(CCA_JSON, "cca_rules.json")
                    st.success("✅ Upload complete!") if ok else st.error("❌ Upload failed")

            with c2:
                st.markdown("### 📥 Drive से Download")
                if st.button("📥 Case Laws Download", use_container_width=True):
                    with st.spinner("Download हो रहा है..."):
                        ok = ds.download_json("case_laws.json", CASE_LAWS_JSON)
                    if ok:
                        load_case_laws.clear()
                        st.success("✅ Download complete!")
                    else:
                        st.error("❌ File not found on Drive")
                if st.button("📥 CCA Rules Download", use_container_width=True):
                    with st.spinner("Download हो रहा है..."):
                        ok = ds.download_json("cca_rules.json", CCA_JSON)
                    if ok:
                        load_cca_data.clear()
                        st.success("✅ Download complete!")
        else:
            st.warning("Drive connect नहीं है।")
            with st.expander("📋 Setup Guide"):
                st.markdown("""
1. Google Cloud Console → Service Account बनाएं
2. Drive API enable करें
3. JSON credentials Streamlit Secrets में डालें
4. Drive folder को Service Account email के साथ share करें
                """)
    except Exception as e:
        st.warning(f"Drive module: {e}")

# ══════════════════════════════════════════════════════
# PAGE: Settings
# ══════════════════════════════════════════════════════
elif page == "⚙️ Settings & API Keys":
    st.header("⚙️ Settings & API Keys")

    tab1, tab2 = st.tabs(["🔑 API Keys", "📂 Data Management"])

    with tab1:
        st.markdown("### API Keys Status")
        for p in AI_PROVIDERS:
            key = p["key_env"]
            status = "✅ Active" if (key and not key.startswith("your_")) else "❌ Not Set"
            col1, col2 = st.columns([2,1])
            col1.markdown(f"**{p['name']}** — `{p['model']}`")
            col2.markdown(status)

        st.divider()
        st.code("""GROQ_API_KEY=gsk_xxxxxxxxxxxx
GEMINI_API_KEY=AIzaSyxxxxxxxxxx
DEEPSEEK_API_KEY=sk-xxxxxxxxxx
OPENROUTER_API_KEY=sk-or-xxxxxxxxxx
GOOGLE_DRIVE_FOLDER_ID=1xxxxxxxxxxxxxxxxxxx""", language="bash")

        st.markdown("""
**Free API Keys:**
- [Groq](https://console.groq.com) — 6000 req/day
- [Gemini](https://aistudio.google.com) — 1500 req/day
- [DeepSeek](https://platform.deepseek.com) — Legal reasoning
- [OpenRouter](https://openrouter.ai) — $1 free credit
        """)

    with tab2:
        st.markdown("### 📂 Data Management")
        bns_data  = load_bns_data()
        case_laws = load_case_laws()
        rti_data  = load_rti_data()
        cca_data  = load_cca_data()

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("BNS Sections", len(bns_data))
        c2.metric("RTI Sections", len(rti_data))
        c3.metric("CCA Rules", len(cca_data))
        c4.metric("Case Laws", len(case_laws))

        if st.button("🔄 Cache Clear करें"):
            load_bns_data.clear()
            load_case_laws.clear()
            load_rti_data.clear()
            load_cca_data.clear()
            st.success("Cache cleared!")

        st.divider()
        st.markdown("### JSON Import")
        up_type = st.selectbox("कौन सा data import करें?",
            ["BNS Sections", "RTI Sections", "CCA Rules", "Case Laws"])
        json_up = st.file_uploader("JSON file upload करें", type=["json"], key="json_imp")
        if json_up:
            try:
                new_data = json.load(json_up)
                path_map = {
                    "BNS Sections": BNS_JSON,
                    "RTI Sections": RTI_JSON,
                    "CCA Rules": CCA_JSON,
                    "Case Laws": CASE_LAWS_JSON
                }
                with open(path_map[up_type], "w", encoding="utf-8") as f:
                    json.dump(new_data, f, ensure_ascii=False, indent=2)
                load_bns_data.clear()
                load_rti_data.clear()
                load_cca_data.clear()
                load_case_laws.clear()
                st.success(f"✅ {len(new_data)} items import हुए!")
            except Exception as e:
                st.error(f"Import error: {e}")

        st.divider()
        st.markdown("### Case Laws Export")
        if case_laws:
            st.download_button("📥 Case Laws Export (JSON)",
                data=json.dumps(case_laws, ensure_ascii=False, indent=2),
                file_name="case_laws_backup.json", mime="application/json")
