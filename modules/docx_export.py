"""
modules/docx_export.py
DOCX Export — जमानत विरोध आख्या और अन्य documents
Hindi Mangal font, proper formatting
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_akhya_docx(akhya_text: str, fir_details: str = "") -> bytes:
    """
    जमानत विरोध आख्या को formatted DOCX में convert करें
    Returns: bytes (for st.download_button)
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

    # ── Header ──────────────────────────────────────────
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("जमानत विरोध आख्या")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Dark red

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_header.add_run("(BAIL OPPOSITION REPORT)")
    sub_run.bold = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # ── FIR Details box ─────────────────────────────────
    if fir_details.strip():
        fir_para = doc.add_paragraph()
        fir_para.style = doc.styles['Normal']
        border_run = fir_para.add_run(f"मुकदमे का विवरण: {fir_details}")
        border_run.bold = True
        border_run.font.size = Pt(11)

        doc.add_paragraph()

    # ── Date line ────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"दिनांक: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_paragraph()

    # ── Main Content ─────────────────────────────────────
    lines = akhya_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # प्रस्तर / heading lines
        if line.startswith("प्रस्तर") or line.startswith("#"):
            line = line.lstrip("#").strip()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("अतः") or line.startswith("इसलिए"):
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = para.add_run(line)
            run.font.size = Pt(11)

    # ── Footer signature ─────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_para.add_run("लोक अभियोजक / विवेचक")
    sig_run.bold = True

    sig_para2 = doc.add_paragraph()
    sig_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para2.add_run("हस्ताक्षर: ____________________")

    sig_para3 = doc.add_paragraph()
    sig_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para3.add_run(f"दिनांक: {datetime.now().strftime('%d/%m/%Y')}")

    # ── Save to bytes ────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def create_section_report_docx(section_data: dict, ai_explanation: str = "") -> bytes:
    """
    किसी धारा का reference card DOCX बनाएं
    """
    doc = Document()

    section_obj = doc.sections[0]
    section_obj.left_margin  = Inches(1.25)
    section_obj.right_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"धारा {section_data.get('bns','—')} — Reference Card")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)

    doc.add_paragraph()

    # Table of info
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ("BNS धारा", section_data.get("bns", "—")),
        ("IPC (पुरानी)", section_data.get("ipc", "—")),
        ("अपराध", section_data.get("offence_hi", section_data.get("offence","—"))),
        ("सज़ा", section_data.get("punishment", "—")),
        ("जमानत स्थिति", "जमानती" if section_data.get("bailable") else "गैर-जमानती"),
    ]

    for i, (k, v) in enumerate(rows_data):
        row = table.rows[i]
        kc = row.cells[0]
        vc = row.cells[1]
        kc.text = k
        vc.text = str(v)
        kc.paragraphs[0].runs[0].bold = True

    if ai_explanation:
        doc.add_paragraph()
        exp_title = doc.add_paragraph()
        exp_title.add_run("AI व्याख्या:").bold = True

        for line in ai_explanation.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
